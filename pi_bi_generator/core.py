"""
PI/BI-Generator Core-Logik
===========================
Reine Datenlogik ohne UI-Abhängigkeiten — liest einen VLB-ONIX-XML-Datensatz
und erzeugt daraus Presse- (PI) und Buchinformation (BI) als docx + html.

Ablauf:
1. lade_buchdaten(xml)          -> Buchdaten (aus ONIX 3.1)
2. lade_cover_datei / _web      -> Cover-Bytes (lokal oder vom Webserver)
3. generiere_docx / generiere_html(vorlage, buch, ...) -> gefüllte Dokumente

Die docx-Vorlagen (vorlagen/*.docx) werden einmalig mit baue_docx_vorlagen()
aus den Muster-Dokumenten in beispiele/ erzeugt und mitgeliefert; jede
buchspezifische Textzeile steckt darin als Platzhalter ({{TITEL}} usw.) in
genau einem Run, damit das Ersetzen zur Laufzeit robust ist.
"""

from __future__ import annotations

import html
import io
import json
import os
import re
import sys
import urllib.request
import xml.etree.ElementTree as ET
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image


# ---------------------------------------------------------------------
# Ordner (Config im "data"-Unterordner, Vorlagen als Ressourcen)
# ---------------------------------------------------------------------

def _base_dir() -> Path:
    """Ordner neben der .exe (PyInstaller-Build) bzw. neben dem Tool-Code."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent


# App-mutierte Laufzeitdaten (config.json) liegen getrennt vom Programm.
APP_DIR = _base_dir() / "data"
APP_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_PFAD = APP_DIR / "config.json"


def _vorlagen_dir() -> Path:
    """Nur-Lese-Vorlagen (docx/html). Im PyInstaller-Build liegen sie unter
    sys._MEIPASS (per datas gebündelt), sonst neben diesem Modul."""
    if getattr(sys, "frozen", False):
        p = Path(getattr(sys, "_MEIPASS", "")) / "pi_bi_generator" / "vorlagen"
        if p.exists():
            return p
        return _base_dir() / "vorlagen"
    return Path(__file__).parent / "vorlagen"


VORLAGEN_DIR = _vorlagen_dir()


# ---------------------------------------------------------------------
# Konfiguration
# ---------------------------------------------------------------------

DEFAULT_CONFIG = {
    "output_dir": "pi_bi_output",
    "last_input_dir": "",
    "verlag_name": "verlag regionalkultur",
    # Fester ISBN-Präfix des Verlags (978-3-95505-…) für Formatierung/Kurzcode.
    "isbn_prefix": "978-3-95505",
    # Basis-URLs für die HTML-Fassung (Cover-Thumbnail + Blick-ins-Buch-PDF).
    # Diese Web-Assets werden vom Nutzer separat hochgeladen.
    "newsletter_base_url": "https://verlag-regionalkultur.de/newsletter_/",
    "bib_base_url": "https://verlag-regionalkultur.de/presse/bib/",
    # Fallback-Link auf die Webshop-Detailseite (nicht aus der XML ableitbar).
    "detail_fallback_url": "https://verlag-regionalkultur.de/",
    # Einband-Code (ONIX b012) -> Klartext. Nur BB ist datenbelegt.
    "einband_map": {
        "BB": "fester Einband",
        "BC": "kartoniert",
        "BE": "Klappenbroschur",
        "PB": "kartoniert",
        "BZ": "Leinen",
    },
    "cover_timeout": 15,
}


def lade_config() -> dict:
    if CONFIG_PFAD.exists():
        with open(CONFIG_PFAD, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        for k, v in DEFAULT_CONFIG.items():
            cfg.setdefault(k, v)
        return cfg
    with open(CONFIG_PFAD, "w", encoding="utf-8") as f:
        json.dump(DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
    return dict(DEFAULT_CONFIG)


def speichere_config(cfg: dict) -> None:
    with open(CONFIG_PFAD, "w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------
# Datenmodell
# ---------------------------------------------------------------------

class Buchdaten:
    __slots__ = (
        "isbn13", "isbn13_formatiert", "shortcode",
        "titel", "serientitel", "band", "band_text",
        "herausgeber", "autoren",
        "editoren_slash", "editoren_und", "mitwirkende",
        "umfang_zeile", "titel_band",
        "verlag", "preis", "verlag_isbn_preis", "datum",
        "cover_url", "werbetext_absaetze",
    )

    def __init__(self, **kw):
        for name in self.__slots__:
            setattr(self, name, kw.get(name, ""))
        if not self.herausgeber:
            self.herausgeber = []
        if not self.autoren:
            self.autoren = []
        if not self.werbetext_absaetze:
            self.werbetext_absaetze = []


# ---------------------------------------------------------------------
# ONIX-Parser (3.1, Kurz-Tags, ohne Namespace)
# ---------------------------------------------------------------------

def _txt(el) -> str:
    return (el.text or "").strip() if el is not None else ""


# Einleitungs-Label, die aus dem Werbetext entfernt werden sollen.
_INTRO_LABELS = ("Aus dem Vorwort",)


def _entferne_intro(absaetze: list[str]) -> list[str]:
    """Entfernt ein führendes Label wie 'Aus dem Vorwort:' — egal ob als
    eigener Absatz oder als Prefix des ersten Absatzes."""
    if not absaetze:
        return absaetze
    for label in _INTRO_LABELS:
        pat = re.compile(rf"^\s*{re.escape(label)}\s*[:\-–—]?\s*", re.IGNORECASE)
        m = pat.match(absaetze[0])
        if m:
            rest = absaetze[0][m.end():].strip()
            return ([rest] if rest else []) + absaetze[1:]
    return absaetze


def _join_und(namen: list[str]) -> str:
    """['A','B','C'] -> 'A, B und C'."""
    namen = [n for n in namen if n]
    if not namen:
        return ""
    if len(namen) == 1:
        return namen[0]
    return ", ".join(namen[:-1]) + " und " + namen[-1]


def _kontributoren(product, rolle: str) -> list[str]:
    dd = product.find("descriptivedetail")
    if dd is None:
        return []
    beitraege = []
    for c in dd.findall("contributor"):
        if _txt(c.find("b035")) == rolle:
            seq = _txt(c.find("b034"))
            name = _txt(c.find("b036"))
            if name:
                beitraege.append((int(seq) if seq.isdigit() else 999, name))
    return [n for _, n in sorted(beitraege)]


def lade_buchdaten(xml_pfad, cfg: dict | None = None) -> Buchdaten:
    cfg = cfg or DEFAULT_CONFIG
    einband_map = cfg.get("einband_map", DEFAULT_CONFIG["einband_map"])
    isbn_prefix = cfg.get("isbn_prefix", DEFAULT_CONFIG["isbn_prefix"])

    root = ET.parse(str(xml_pfad)).getroot()
    product = root.find("product")
    if product is None:
        raise ValueError("Kein <product>-Element in der XML gefunden — "
                         "ist das eine VLB-ONIX-Datei?")
    dd = product.find("descriptivedetail")

    # -- ISBN / Kurzcode ------------------------------------------------
    isbn13 = ""
    for pid in product.findall("productidentifier"):
        if _txt(pid.find("b221")) in ("03", "15"):
            isbn13 = _txt(pid.find("b244"))
            break
    if not isbn13 or len(isbn13) != 13 or not isbn13.isdigit():
        raise ValueError(f"Keine gültige ISBN-13 gefunden (gelesen: {isbn13!r}).")
    e = isbn13
    isbn13_formatiert = f"{isbn_prefix}-{e[9:12]}-{e[12]}"
    shortcode = f"{e[7:9]}-{e[9:12]}-{e[12]}"

    # -- Titel (Produktebene, NICHT die Collection-Kopie) ---------------
    titel = ""
    if dd is not None:
        for td in dd.findall("titledetail"):
            if _txt(td.find("b202")) == "01":
                for te in td.findall("titleelement"):
                    if _txt(te.find("x409")) == "01":
                        titel = _txt(te.find("b203"))
                        break
            if titel:
                break

    # -- Serie / Band ---------------------------------------------------
    serientitel = ""
    band = ""
    if dd is not None:
        coll = dd.find("collection")
        if coll is not None:
            for te in coll.findall("titledetail/titleelement"):
                lvl = _txt(te.find("x409"))
                if lvl == "02" and not serientitel:
                    serientitel = _txt(te.find("b203"))
                if lvl == "01" and not band:
                    band = _txt(te.find("x410"))
    band_text = f"Band {band}" if band else ""

    # -- Beteiligte -----------------------------------------------------
    herausgeber = _kontributoren(product, "B01")
    autoren = _kontributoren(product, "A01")
    editoren_slash = (" / ".join(herausgeber) + " (Hrsg.)") if herausgeber else ""
    editoren_und = (_join_und(herausgeber) + " (Hrsg.)") if herausgeber else ""
    mitwirkende = (f"Mit Beiträgen von {_join_und(autoren)}."
                   if autoren else "")

    # -- Umfang / Einband ----------------------------------------------
    seiten = _txt(dd.find("extent/b219")) if dd is not None else ""
    anzahl_abb = _txt(dd.find("b125")) if dd is not None else ""
    abb_text = _txt(dd.find("b062")) if dd is not None else ""
    einband_code = _txt(dd.find("b012")) if dd is not None else ""
    einband = einband_map.get(einband_code, "")
    if not einband and einband_code:
        einband = einband_code  # unbekannter Code: roh anzeigen
        print(f"⚠ Unbekannter Einband-Code b012={einband_code!r} — "
              f"bitte einband_map in config.json ergänzen.")
    teile = []
    if seiten:
        teile.append(f"{seiten} Seiten")
    if anzahl_abb and abb_text:
        teile.append(f"mit {anzahl_abb} {abb_text}")
    umfang_zeile = " ".join(teile)
    if einband:
        umfang_zeile = (umfang_zeile + ", " if umfang_zeile else "") + einband
    if umfang_zeile:
        umfang_zeile += "."

    titel_band = titel + (f". {band_text}." if band_text else ".")

    # -- Verlag / Preis / Datum ----------------------------------------
    verlag = _txt(product.find("publishingdetail/publisher/b081")) \
        or cfg.get("verlag_name", "")
    preis = ""
    for price in product.findall("productsupply/supplydetail/price"):
        terr = price.find("territory")
        x449 = _txt(terr.find("x449")) if terr is not None else ""
        if x449 == "DE":
            betrag = _txt(price.find("j151"))
            waehrung = _txt(price.find("j152")) or "EUR"
            if betrag:
                try:
                    preis = f"{waehrung} {float(betrag):.2f}".replace(".", ",")
                except ValueError:
                    preis = f"{waehrung} {betrag}"
            break
    verlag_isbn_preis = f"{verlag}, ISBN {isbn13_formatiert}."
    if preis:
        verlag_isbn_preis += f" {preis}."

    roh_datum = _txt(product.find("publishingdetail/publishingdate/b306"))
    datum = ""
    if len(roh_datum) == 8 and roh_datum.isdigit():
        datum = f"{roh_datum[6:8]}.{roh_datum[4:6]}.{roh_datum[0:4]}"

    # -- Cover-URL (Vorderseite x436=01) -------------------------------
    cover_url = ""
    for sr in product.findall("collateraldetail/supportingresource"):
        if _txt(sr.find("x436")) == "01":
            cover_url = _txt(sr.find("resourceversion/x435"))
            break

    # -- Werbetext (d104, mehrzeilig) ----------------------------------
    d104 = _txt(product.find("collateraldetail/textcontent/d104"))
    werbetext_absaetze = [ln.strip() for ln in d104.split("\n") if ln.strip()]
    werbetext_absaetze = _entferne_intro(werbetext_absaetze)

    return Buchdaten(
        isbn13=isbn13, isbn13_formatiert=isbn13_formatiert, shortcode=shortcode,
        titel=titel, serientitel=serientitel, band=band, band_text=band_text,
        herausgeber=herausgeber, autoren=autoren,
        editoren_slash=editoren_slash, editoren_und=editoren_und,
        mitwirkende=mitwirkende, umfang_zeile=umfang_zeile, titel_band=titel_band,
        verlag=verlag, preis=preis, verlag_isbn_preis=verlag_isbn_preis,
        datum=datum, cover_url=cover_url, werbetext_absaetze=werbetext_absaetze,
    )


# ---------------------------------------------------------------------
# Cover laden (lokal oder vom Webserver)
# ---------------------------------------------------------------------

def lade_cover_datei(pfad) -> bytes:
    return Path(pfad).read_bytes()


def web_cover_url(cfg: dict, buch: Buchdaten) -> str:
    return f"{cfg.get('newsletter_base_url', '')}{buch.shortcode}.png"


def lade_cover_web(url: str, timeout: float = 15.0) -> bytes:
    req = urllib.request.Request(
        url, headers={"User-Agent": "PiBiGenerator/1.0"})
    with urllib.request.urlopen(req, timeout=timeout) as r:
        return r.read()


# ---------------------------------------------------------------------
# docx-Platzhalter (buchspezifische Zeilen)
# ---------------------------------------------------------------------

TOKEN_WERBETEXT = "{{WERBETEXT}}"


def _mapping(buch: Buchdaten) -> dict:
    """Einzeilige Platzhalter -> Wert (Werbetext separat, siehe unten)."""
    return {
        "{{EDITOREN}}": buch.editoren_slash,
        "{{TITEL}}": buch.titel,
        "{{BAND}}": buch.band_text,
        "{{EDITOREN_UND}}": buch.editoren_und,
        "{{TITEL_BAND}}": buch.titel_band,
        "{{MITWIRKENDE}}": buch.mitwirkende,
        "{{UMFANG_ZEILE}}": buch.umfang_zeile,
        "{{VERLAG_ISBN_PREIS}}": buch.verlag_isbn_preis,
    }


def _textruns(para):
    """Runs eines Absatzes ohne Zeichnungs-Runs (Cover-Bild etc.)."""
    return [r for r in para.runs if r._r.find(qn("w:drawing")) is None]


def _ersetze_in_absaetzen(absaetze, mapping):
    for p in absaetze:
        runs = _textruns(p)
        if not runs:
            continue
        full = "".join(r.text for r in runs)
        if "{{" not in full:
            continue
        neu = full
        for k, v in mapping.items():
            neu = neu.replace(k, v)
        if neu != full:
            runs[0].text = neu
            for r in runs[1:]:
                r.text = ""


def _alle_absaetze(doc):
    """Alle Absätze aus Body + Kopf-/Fußzeilen + Tabellen."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for sec in doc.sections:
        for hf in (sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer):
            if hf is not None:
                yield from hf.paragraphs


def _ersetze_platzhalter(doc, mapping):
    _ersetze_in_absaetzen(list(_alle_absaetze(doc)), mapping)


def _fuelle_werbetext(doc, absaetze: list[str]):
    """Ersetzt den {{WERBETEXT}}-Absatz durch je einen Absatz pro Zeile,
    unter Beibehaltung von Formatierung/Absatzstil der Vorlage."""
    ziel = None
    for p in doc.paragraphs:
        if p.text.strip() == TOKEN_WERBETEXT:
            ziel = p
            break
    if ziel is None:
        return
    for line in (absaetze or [""]):
        neu = deepcopy(ziel._p)
        para = Paragraph(neu, ziel._parent)
        runs = _textruns(para)
        if runs:
            runs[0].text = line
            for r in runs[1:]:
                r.text = ""
        ziel._p.addprevious(neu)
    ziel._p.getparent().remove(ziel._p)


# ---------------------------------------------------------------------
# Cover-Bild im docx austauschen
# ---------------------------------------------------------------------

def _cover_zeichnung(doc):
    """Das <w:drawing>, das ein eingebettetes Bild (a:blip) referenziert."""
    for p in doc.paragraphs:
        for r in p.runs:
            drawing = r._r.find(qn("w:drawing"))
            if drawing is not None and drawing.find(".//" + qn("a:blip")) is not None:
                return drawing
    return None


def _setze_cover_groesse(drawing, breite_beibehalten: bool, w: int, h: int):
    """Passt die Bildbox an das Seitenverhältnis des neuen Covers an
    (Breite bleibt, Höhe wird nachgezogen), damit nichts verzerrt wird."""
    if not w or not h:
        return
    ext = drawing.find(".//" + qn("wp:extent"))
    if ext is None or not ext.get("cx"):
        return
    cx = int(ext.get("cx"))
    cy = round(cx * h / w)
    ext.set("cy", str(cy))
    for aext in drawing.findall(".//" + qn("a:ext")):
        if aext.get("cx"):  # nur die größenrelevanten a:ext (nicht extLst-URI)
            aext.set("cx", str(cx))
            aext.set("cy", str(cy))


def _tausche_cover(doc, cover_bytes: bytes) -> bool:
    """Überschreibt das Cover-Bild (word/media/image1.jpeg) und zieht die
    Bildbox auf das echte Seitenverhältnis. Gibt True bei Erfolg zurück."""
    cover_rel = None
    for rel in doc.part.rels.values():
        if "image" in rel.reltype and rel.target_ref.endswith("image1.jpeg"):
            cover_rel = rel
            break
    if cover_rel is None:
        return False
    cover_rel.target_part._blob = cover_bytes
    try:
        im = Image.open(io.BytesIO(cover_bytes))
        drawing = _cover_zeichnung(doc)
        if drawing is not None:
            _setze_cover_groesse(drawing, True, im.size[0], im.size[1])
    except Exception as ex:  # Größenanpassung ist optional
        print(f"⚠ Cover-Größe konnte nicht angepasst werden: {ex}")
    return True


# ---------------------------------------------------------------------
# docx erzeugen
# ---------------------------------------------------------------------

def generiere_docx(vorlage_pfad, buch: Buchdaten, cover_bytes: bytes | None,
                   ziel_pfad) -> None:
    doc = Document(str(vorlage_pfad))
    _fuelle_werbetext(doc, buch.werbetext_absaetze)
    _ersetze_platzhalter(doc, _mapping(buch))
    if cover_bytes:
        _tausche_cover(doc, cover_bytes)
    ziel_pfad = Path(ziel_pfad)
    ziel_pfad.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ziel_pfad))


# ---------------------------------------------------------------------
# html erzeugen
# ---------------------------------------------------------------------

def _html_mapping(buch: Buchdaten, detail_url: str, cfg: dict) -> dict:
    nl = cfg.get("newsletter_base_url", "")
    bib = cfg.get("bib_base_url", "")
    werbetext_html = "<br />\n".join(
        html.escape(a) for a in buch.werbetext_absaetze)
    return {
        "{{EDITOREN}}": html.escape(buch.editoren_slash),
        "{{TITEL_HTML}}": html.escape(buch.titel),
        "{{BAND}}": html.escape(buch.band),
        "{{BAND_TEXT}}": html.escape(buch.band_text),
        "{{DATUM}}": html.escape(buch.datum),
        "{{DETAIL_URL}}": html.escape(detail_url, quote=True),
        "{{COVER_THUMB_URL}}": html.escape(f"{nl}{buch.shortcode}.png", quote=True),
        "{{BIB_URL}}": html.escape(f"{bib}bib_{buch.shortcode}.pdf", quote=True),
        "{{WERBETEXT_HTML}}": werbetext_html,
        "{{UMFANG_ZEILE}}": html.escape(buch.umfang_zeile),
        "{{VERLAG_ISBN_PREIS}}": html.escape(buch.verlag_isbn_preis),
        "{{TITEL}}": html.escape(buch.titel),
    }


def generiere_html(vorlage_pfad, buch: Buchdaten, detail_url: str, cfg: dict,
                   ziel_pfad) -> None:
    text = Path(vorlage_pfad).read_text(encoding="utf-8")
    for k, v in _html_mapping(buch, detail_url, cfg).items():
        text = text.replace(k, v)
    ziel_pfad = Path(ziel_pfad)
    ziel_pfad.parent.mkdir(parents=True, exist_ok=True)
    ziel_pfad.write_text(text, encoding="utf-8")


# ---------------------------------------------------------------------
# Vorlagen-Erzeugung (Entwickler-Routine, einmalig)
# ---------------------------------------------------------------------

# Absatz-Index -> Platzhalter im Muster-docx (PI und BI identisch bis P16).
_VORLAGE_TOKENS = {
    1: "{{EDITOREN}}",
    3: "{{TITEL}}",
    5: "{{BAND}}",              # nur Text-Run; Cover-Zeichnung bleibt
    7: TOKEN_WERBETEXT,         # P8, P9 werden gelöscht
    12: "{{EDITOREN_UND}}",
    13: "{{TITEL_BAND}}",
    14: "{{MITWIRKENDE}}",
    15: "{{UMFANG_ZEILE}}",
    16: "{{VERLAG_ISBN_PREIS}}",
}


def _setze_token(para, token: str):
    runs = _textruns(para)
    if not runs:
        return
    runs[0].text = token
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


def baue_docx_vorlage(muster_pfad, ziel_pfad) -> None:
    doc = Document(str(muster_pfad))
    paras = doc.paragraphs

    # Plausibilitätsprüfung: erwartete Ankertexte
    if "Hrsg." not in paras[1].text or "ISBN" not in paras[16].text:
        raise ValueError(f"Unerwartete Absatzstruktur in {muster_pfad} — "
                         f"Vorlagen-Mapping passt nicht.")

    # Blurb-Folgeabsätze (P8, P9) merken und später entfernen
    zu_loeschen = [paras[8]._p, paras[9]._p]
    for idx, token in _VORLAGE_TOKENS.items():
        _setze_token(paras[idx], token)
    for p in zu_loeschen:
        p.getparent().remove(p)

    # Cover-Zuschnitt (srcRect) entfernen -> Laufzeit setzt echtes Verhältnis
    drawing = _cover_zeichnung(doc)
    if drawing is not None:
        for sr in drawing.findall(".//" + qn("a:srcRect")):
            sr.getparent().remove(sr)

    ziel_pfad = Path(ziel_pfad)
    ziel_pfad.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ziel_pfad))


def baue_docx_vorlagen(beispiel_dir=None, ziel_dir=None) -> None:
    """Erzeugt vorlagen/pi_vorlage.docx und bi_vorlage.docx aus den Mustern."""
    beispiel_dir = Path(beispiel_dir or (_base_dir() / "beispiele"
                        / "Bausteine Bruchsal, Bd. 5"))
    ziel_dir = Path(ziel_dir or VORLAGEN_DIR)
    baue_docx_vorlage(beispiel_dir / "PI_05-559-2.docx",
                      ziel_dir / "pi_vorlage.docx")
    baue_docx_vorlage(beispiel_dir / "BI_05-559-2.docx",
                      ziel_dir / "bi_vorlage.docx")
    print(f"Vorlagen geschrieben nach {ziel_dir}")
