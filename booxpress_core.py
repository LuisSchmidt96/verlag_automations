"""
BOOXpress-Etiketten Core-Logik
================================
Reine Datenlogik ohne UI-Abhängigkeiten — wird sowohl von der GUI-App
als auch von einem ggf. zukünftigen CLI-Skript verwendet.
"""

from __future__ import annotations

import io
import json
import re
import sys
from pathlib import Path

import pandas as pd
from barcode import ITF
from barcode.writer import ImageWriter
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# ---------------------------------------------------------------------
# App-Ordner (für Config, Counter, Komm-Liste)
# ---------------------------------------------------------------------

def app_dir() -> Path:
    """Ordner neben der .exe (bei PyInstaller-Build) bzw. neben dem Skript."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent


APP_DIR = app_dir()
CONFIG_PFAD = APP_DIR / "config.json"
PAKETNR_PFAD = APP_DIR / "paketnr.txt"


# ---------------------------------------------------------------------
# Konfiguration
# ---------------------------------------------------------------------

DEFAULT_CONFIG = {
    "verlag_knr_booxpress": "068249",
    "kommliste_pfad": "kommliste.xlsx",
    "absender_name": "verlag regionalkultur",
    "absender_adresse": "Bahnhofstr. 2, 76698 Ubstadt-Weiher",
    "libri_pattern": "^Libri",
    "libri_strasse": "Europaallee 1",
    "libri_plz_ort": "36244 Bad Hersfeld",
    "output_dir": "etiketten_output",
    "last_input_dir": "",
}


def lade_config() -> dict:
    if CONFIG_PFAD.exists():
        with open(CONFIG_PFAD, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        for k, v in DEFAULT_CONFIG.items():
            cfg.setdefault(k, v)
        return cfg
    with open(CONFIG_PFAD, "w", encoding="utf-8") as f:
        json.dump(DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
    return dict(DEFAULT_CONFIG)


def speichere_config(cfg: dict) -> None:
    with open(CONFIG_PFAD, "w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------
# Paketnummer-Counter
# ---------------------------------------------------------------------

def lade_paketnr() -> int:
    if not PAKETNR_PFAD.exists():
        PAKETNR_PFAD.write_text("1\n", encoding="utf-8")
        return 1
    try:
        return int(PAKETNR_PFAD.read_text(encoding="utf-8").strip())
    except (ValueError, OSError):
        return 1


def speichere_paketnr(n: int) -> None:
    PAKETNR_PFAD.write_text(f"{n}\n", encoding="utf-8")


# ---------------------------------------------------------------------
# Datenmodell + Lade-Funktionen
# ---------------------------------------------------------------------

class Etikett:
    __slots__ = ("name", "strasse", "plz_ort", "barcode_data", "belegnr")

    def __init__(self, name, strasse, plz_ort, barcode_data, belegnr):
        self.name = name
        self.strasse = strasse
        self.plz_ort = plz_ort
        self.barcode_data = barcode_data
        self.belegnr = belegnr


def lade_auftraege(pfad: Path) -> pd.DataFrame:
    """Lexware-Export: Zeile 1 = Titel ('Aufträge Verkauf ...'), Zeile 2 = Header."""
    df = pd.read_excel(pfad, dtype=str, header=1)
    df.columns = [c.strip() for c in df.columns]
    return df


def lade_kommliste(pfad: Path) -> pd.DataFrame:
    df = pd.read_excel(pfad, dtype=str)
    df.columns = [c.strip() for c in df.columns]
    if "Kommittent" in df.columns:
        df = df[df["Kommittent"] == "1"].copy()
    df = df[df["VD"].notna() & (df["VD"] != "")].copy()
    return df


def _strip_leading_one(kd_nr: str) -> str:
    s = str(kd_nr).strip()
    return s[1:] if s.startswith("1") else s


def _clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    s = str(value).strip()
    return "" if s.lower() == "nan" else s


# ---------------------------------------------------------------------
# Matching & Etikett-Erzeugung
# ---------------------------------------------------------------------

def generiere_etiketten(auftraege, kommliste, cfg, start_paketnr):
    bx_by_vd = {str(row["VD"]).strip(): row for _, row in kommliste.iterrows()}
    verlag_knr = cfg["verlag_knr_booxpress"]
    libri_re = re.compile(cfg["libri_pattern"], re.IGNORECASE)

    etiketten, warnungen = [], []
    paketnr = start_paketnr

    for _, auftrag in auftraege.iterrows():
        kd_nr = _clean(auftrag.get("Kd.-Nr.", ""))
        if not kd_nr:
            continue
        matchcode = _clean(auftrag.get("Matchcode", ""))
        belegnr = _clean(auftrag.get("Belegnr.", ""))

        if not kd_nr.startswith("1"):
            warnungen.append({"belegnr": belegnr, "kd_nr": kd_nr,
                              "matchcode": matchcode, "grund": "Endkunde — kein BOOXpress"})
            continue

        vd = _strip_leading_one(kd_nr)
        bx_row = bx_by_vd.get(vd)
        if bx_row is None:
            warnungen.append({"belegnr": belegnr, "kd_nr": kd_nr,
                              "matchcode": matchcode,
                              "grund": f"VD {vd} nicht in Komm-Liste (vmtl. GLS)"})
            continue

        name1 = _clean(bx_row.get("Name1"))
        name2 = _clean(bx_row.get("Name2"))
        strasse = _clean(bx_row.get("Strasse"))
        hausnr = _clean(bx_row.get("Hausnummer"))
        plz = _clean(bx_row.get("PLZ"))
        ort = _clean(bx_row.get("Ort"))

        if libri_re.match(name1):
            empf_name = f"Libri_{vd}"
            empf_strasse = cfg["libri_strasse"]
            empf_plz_ort = cfg["libri_plz_ort"]
        else:
            empf_name = name1 + (f"\n{name2}" if name2 else "")
            empf_strasse = f"{strasse} {hausnr}".strip()
            empf_plz_ort = f"{plz} {ort}".strip()

        barcode_data = f"{verlag_knr}{paketnr:06d}{vd.zfill(6)}"
        assert len(barcode_data) == 18 and barcode_data.isdigit()

        etiketten.append(Etikett(empf_name, empf_strasse, empf_plz_ort,
                                  barcode_data, belegnr))
        paketnr += 1

    return etiketten, warnungen


# ---------------------------------------------------------------------
# Barcode
# ---------------------------------------------------------------------

def barcode_png_bytes(data: str) -> bytes:
    """ITF (Interleaved 2 of 5) — Optionen sind so abgestimmt, dass das
    natürliche PNG-Seitenverhältnis ungefähr 9.05cm:2.117cm = 4.275 trifft,
    damit die Anzeige in der docx-Zelle nicht verzerrt wird.

    text_distance muss > pt2mm(font_size) sein, sonst überlappt der lesbare
    Text die Striche (Bug der ImageWriter-Mittel-Anker-Darstellung).
    """
    buf = io.BytesIO()
    ITF(data, writer=ImageWriter()).write(buf, options={
        "module_width": 0.32,
        "module_height": 14.0,
        "quiet_zone": 1.0,
        "font_size": 11,
        "text_distance": 5.0,
        "write_text": True,
    })
    return buf.getvalue()


# ---------------------------------------------------------------------
# docx-Aufbau — Layout 1:1 zum Original-Muster (Etiketten_BOOXpress.docx)
#
# Specs (per XML aus dem Muster extrahiert):
#   A4 21×29.7cm, Margins: top=0.85cm (482 dxa), sonst 0
#   2-spaltige Tabelle, Gesamtbreite 21cm (11906 dxa), fixed layout
#   Zellbreite je 10.5cm (5953 dxa), Zellinnenränder 15 dxa li/re
#   Zeilenhöhe EXAKT 7cm (3969 dxa), cantSplit
#   Zelle vertikal zentriert
#
#   Zellinhalt (Reihenfolge!):
#     1) "BOOXpress" 18pt fett, zentriert  (sz=36 half-points)
#     2) Barcode-Bild zentriert, 9.05 × 2.117 cm
#     3) Mini-Leerabsatz (Mark-Größe 6pt = sz 12)
#     4) "verlag regionalkultur, "  11pt fett   (tab 540)
#     5) "Bahnhofstr. 2, 76698 Ubstadt-Weiher" 11pt fett UNTERSTRICHEN (tab 540)
#     6) "Empfänger:"  11pt NICHT fett        (tab 527)
#     7..) Empfänger-Name(n) / Strasse / PLZ Ort  11pt fett  (tab 527)
#
#   Alle Absätze haben Einzug links = 180 dxa (~0.32 cm).
#   Bei Zeilenhöhe 7cm passen 4 Zeilen auf A4 → 8 Etiketten pro Seite.
# ---------------------------------------------------------------------

ZEILENHOEHE_DXA = 3969       # 7.0 cm
SPALTENBREITE_DXA = 5953     # 10.5 cm
TABELLE_BREITE_DXA = 11906   # 21 cm
ZELLRAND_DXA = 15            # Innenrand links/rechts
EINZUG_DXA = 180             # 0.32 cm
TABPOS_SENDER = 540
TABPOS_EMPF = 527
TABPOS_SEKUNDAER = 2685
BARCODE_BREITE_CM = 9.05
BARCODE_HOEHE_CM = 2.117

SZ_TITEL = 36     # 18pt
SZ_BODY = 22      # 11pt
SZ_SPACER = 12    # 6pt (nur Absatzmarken-Größe für Mini-Spacer)


ETIKETTEN_PRO_SEITE = 8


def baue_docx(etiketten, output_pfad: Path, cfg: dict,
              start_position: int = 1) -> None:
    """start_position: 1..8 (Pos. auf der ersten Seite, von oben links nach
    unten rechts gezählt). Erlaubt das Weiterverwenden teilbedruckter Bögen.
    Vorausgegangene Positionen werden als leere Zellen erzeugt."""
    if not 1 <= start_position <= ETIKETTEN_PRO_SEITE:
        raise ValueError(
            f"start_position muss 1..{ETIKETTEN_PRO_SEITE} sein, war {start_position}"
        )

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0.85)
    sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0)
    sec.right_margin = Cm(0)

    etiketten = [None] * (start_position - 1) + list(etiketten)
    if len(etiketten) % 2 == 1:
        etiketten = etiketten + [None]

    zeilen = len(etiketten) // 2
    table = doc.add_table(rows=zeilen, cols=2)
    table.autofit = False
    _tabelle_konfig(table)

    for ri in range(zeilen):
        row = table.rows[ri]
        _zeile_konfig(row)
        for ci in range(2):
            cell = row.cells[ci]
            _zelle_konfig(cell)
            _zelle_fuellen(cell, etiketten[ri * 2 + ci], cfg)

    output_pfad.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_pfad)


# ---- Tabellen-/Zellen-Setup ----------------------------------------

def _tabelle_konfig(table) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(TABELLE_BREITE_DXA))
    tblW.set(qn("w:type"), "dxa")

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for old in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(old)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ("left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(ZELLRAND_DXA))
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)

    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)

    grid = table._element.find(qn("w:tblGrid"))
    if grid is not None:
        for col in grid.findall(qn("w:gridCol")):
            col.set(qn("w:w"), str(SPALTENBREITE_DXA))


def _zeile_konfig(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(ZEILENHOEHE_DXA))
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def _zelle_konfig(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(SPALTENBREITE_DXA))
    tcW.set(qn("w:type"), "dxa")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ---- Zellinhalt -----------------------------------------------------

def _zelle_fuellen(cell, et, cfg):
    p_first = cell.paragraphs[0]
    if et is None:
        return

    p_first.paragraph_format.space_before = Pt(0)
    p_first.paragraph_format.space_after = Pt(0)
    _setze_pPr(p_first, jc="center", tabs=(TABPOS_SEKUNDAER,))
    r = p_first.add_run("BOOXpress")
    r.bold = True
    r.font.size = Pt(SZ_TITEL / 2)

    _absatz_barcode(cell, et.barcode_data)
    _absatz_spacer(cell)

    _absatz_tabtext(cell, f"{cfg['absender_name']}, ",
                    bold=True, tabs=(TABPOS_SENDER, TABPOS_SEKUNDAER))
    _absatz_tabtext(cell, cfg["absender_adresse"], bold=True, underline=True,
                    tabs=(TABPOS_SENDER, TABPOS_SEKUNDAER))
    _absatz_tabtext(cell, "Empfänger:", bold=False,
                    tabs=(TABPOS_EMPF, TABPOS_SEKUNDAER))

    for line in et.name.split("\n"):
        _absatz_tabtext(cell, line, bold=True, tabs=(TABPOS_EMPF,))
    _absatz_tabtext(cell, et.strasse, bold=True, tabs=(TABPOS_EMPF,))
    _absatz_tabtext(cell, et.plz_ort, bold=True, tabs=(TABPOS_EMPF,))


def _absatz_barcode(cell, data: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _setze_pPr(p, jc="center", tabs=(TABPOS_SEKUNDAER,))
    img = barcode_png_bytes(data)
    p.add_run().add_picture(io.BytesIO(img),
                            width=Cm(BARCODE_BREITE_CM),
                            height=Cm(BARCODE_HOEHE_CM))
    _entferne_bildrahmen(p)


def _entferne_bildrahmen(p) -> None:
    """Setzt explizit <a:noFill/> + <a:ln><a:noFill/></a:ln> auf das pic:spPr,
    damit Word/LibreOffice keinen Default-Rahmen um das Bild zeichnet."""
    spPr = p._p.find(f".//{qn('pic:spPr')}")
    if spPr is None:
        return
    for tag in ("a:noFill", "a:solidFill", "a:gradFill", "a:pattFill", "a:ln"):
        for el in spPr.findall(qn(tag)):
            spPr.remove(el)
    spPr.append(OxmlElement("a:noFill"))
    ln = OxmlElement("a:ln")
    ln.append(OxmlElement("a:noFill"))
    spPr.append(ln)


def _absatz_spacer(cell) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _setze_pPr(p, tabs=(1080, TABPOS_SEKUNDAER), mark_sz=SZ_SPACER)


def _absatz_tabtext(cell, text, *, bold, tabs, underline=False) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _setze_pPr(p, tabs=tabs)

    r_tab = p.add_run()
    r_tab.font.size = Pt(SZ_BODY / 2)
    r_tab._r.append(OxmlElement("w:tab"))

    r_txt = p.add_run(text)
    r_txt.font.size = Pt(SZ_BODY / 2)
    if bold:
        r_txt.bold = True
    if underline:
        r_txt.font.underline = True


# ---- pPr-Helfer (Tabs, Einzug, Justification, Mark-Größe) ----------

def _setze_pPr(p, *, jc=None, tabs=(), left_indent=EINZUG_DXA, mark_sz=None):
    pPr = p._p.get_or_add_pPr()

    for old in pPr.findall(qn("w:tabs")):
        pPr.remove(old)
    if tabs:
        tabs_el = OxmlElement("w:tabs")
        for pos in tabs:
            t = OxmlElement("w:tab")
            t.set(qn("w:val"), "left")
            t.set(qn("w:pos"), str(pos))
            tabs_el.append(t)
        pPr.insert(0, tabs_el)

    if left_indent is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), str(left_indent))

    if jc is not None:
        jc_el = pPr.find(qn("w:jc"))
        if jc_el is None:
            jc_el = OxmlElement("w:jc")
            pPr.append(jc_el)
        jc_el.set(qn("w:val"), jc)

    if mark_sz is not None:
        for old in pPr.findall(qn("w:rPr")):
            pPr.remove(old)
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(mark_sz))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(mark_sz))
        rPr.append(szCs)
        pPr.append(rPr)
